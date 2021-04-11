#!/usr/bin/env python
"""
Sync a static collection from onprem to cloud
"""
from __future__ import print_function

import argparse
import datetime
import email.mime.base
import email.mime.multipart
import email.mime.text
import io
import json
import locale
import logging
import os
import platform
import smtplib
import subprocess
import sys
import time
import xml.etree.cElementTree as ET
from collections import OrderedDict, namedtuple
from pprint import pprint

import docx
import requests
import requests_cache
from docx.enum.style import WD_STYLE_TYPE

try:
    basestring
except NameError:
    basestring = str


TODO = [
    [
        "fix all the REPLACE bits by reading a configuration file",
    ],
    [
        "convert from generating instructions to doing stuff",
    ],
    [
        "convert to proper package",
    ],
    [
        "use JIRA API to autoguess action number based on shortname in the tile",
        "too many choices on <https://pypi.org/search/?q=jira>",
    ],
]


requests_cache.install_cache(
    "test_cache." + ("ascii" if bytes == str else "unicode"),
    backend="sqlite",
    expire_after=60 * 30,
)


HOSTNAME_FULL = platform.node()
HOSTNAME_SHORT = HOSTNAME_FULL.split(".", 1)[0]
SENDER = os.environ["LOGNAME"]
RECIPIENT = SENDER
NOW = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
TODAY = datetime.datetime.now().strftime("%Y-%m-%d")

DEFAULT_CONFIG_TREE_NAME = "**REPLACE(DAAC_CLOUD_ALL_CONFIGREPO_NAME)**"
DEFAULT_CONFIG_TREE_REMOTE = (
    "https://**REPLACE(DAAC_CLOUD_ALL_CONFIGREPO_HOST)**/**REPLACE(DAAC_CLOUD_ALL_CONFIGREPO_ORG)**/"
    + DEFAULT_CONFIG_TREE_NAME
    + ".git"
)
DEFAULT_TARGET_NAME = "**REPLACE(DAAC_CLOUD_LIVE_ENV_NAME)**"


DeploymentTarget = namedtuple(
    "DeploymentTarget",
    ("name", "prefix", "cmr", "cmr_preferred_providers", "search", "dashboard"),
)

KNOWN_TARGETS = (
    DeploymentTarget(
        name="**REPLACE(DAAC_CLOUD_LIVE_ENV_NAME)**",
        prefix="**REPLACE(DAAC_CLOUD_LIVE_CUMULUS_PREFIX)**",
        cmr="https://cmr.earthdata.nasa.gov",
        cmr_preferred_providers=[
            "**REPLACE(DAAC_CLOUD_LIVE_CMR_NAME_NEW)**",
            "**REPLACE(DAAC_CLOUD_LIVE_CMR_NAME_OLD)**",
        ],
        search="https://search.earthdata.nasa.gov",
        dashboard="https://**REPLACE(DAAC_CLOUD_LIVE_CUMULUS_DASHBOARD)**",
    ),
    DeploymentTarget(
        name="**REPLACE(DAAC_CLOUD_STAGING_ENV_NAME)**",
        prefix="**REPLACE(DAAC_CLOUD_STAGING_CUMULUS_PREFIX)**",
        cmr="https://cmr.uat.earthdata.nasa.gov",
        cmr_preferred_providers=[
            "**REPLACE(DAAC_CLOUD_STAGING_CMR_NAME_NEW)**",
            "**REPLACE(DAAC_CLOUD_STAGING_CMR_NAME_OLD)**",
        ],
        search="https://search.uat.earthdata.nasa.gov",
        dashboard="https://**REPLACE(DAAC_CLOUD_STAGING_CUMULUS_DASHBOARD)**",
    ),
)
KNOWN_TARGETS_BY_NAME = {_.name: _ for _ in KNOWN_TARGETS}


def get_target_by_name(target_name):
    return KNOWN_TARGETS_BY_NAME[target_name]


PARSER_OPTIONS = dict(
    _treename=DEFAULT_CONFIG_TREE_NAME,
    _remote=DEFAULT_CONFIG_TREE_REMOTE,
    description="Generate instructions for adding a static dataset to **REPLACE(DAAC_ABBREV)**",
)


class OnPremConfig:
    def __init__(self, target_name=DEFAULT_TARGET_NAME):
        if target_name in (
            "**REPLACE(DAAC_CLOUD_STAGING_ENV_NAME)**",
            "**REPLACE(DAAC_CLOUD_LIVE_ENV_NAME)**",
        ):
            self.env = "**REPLACE(DAAC_ONPREM_LIVE_DATA_ENV_ABBREV)**"
        else:
            raise ValueError("unknown target name", target_name=target_name)
        self.source = (
            "/**REPLACE(DAAC_ONPREM_ALL_DATA_ROOT)**/"
            + self.env
            + "/**REPLACE(DAAC_ONPREM_ALL_DATA_RELPATH)**"
        )


class InstructionsWriter:
    def title(self, text):
        pass

    def heading(self, text, level):
        pass

    def unordered_list_item(self, text):
        pass

    def close_list(self):
        pass

    def step(
        self, description, commands, skip_session_placeholder=False, ignore_error=False
    ):
        pass


class Instructions:
    def __init__(self, collection, jira_issue_key, target):
        self.collection = collection
        self.jira_issue_key = jira_issue_key
        self.target = target

    def prepare(self):
        # TODO fetch issue summary, cf https://developer.atlassian.com/server/jira/platform/rest-apis/
        # jira_issue_rest_url = 'https://bugs.earthdata.nasa.gov/rest/api/latest/issue/'+jira_issue_key
        # jira_issue_details = ...
        # jira_issue_summary = jira_issue_details.fields.summary
        self.jira_issue_url = (
            "https://bugs.earthdata.nasa.gov/browse/" + self.jira_issue_key
        )
        self.jira_issue_summary = self.collection.name + ": Copy to " + self.target.name
        self.title = self.jira_issue_key + " " + self.jira_issue_summary
        self.onprem_details_url = (
            "**REPLACE(DAAC_ONPREM_ALL_INFO_URL)**/details/" + self.collection.name
        )
        self.onprem_download_url = (
            "**REPLACE(DAAC_ONPREM_ALL_DATA_DOWNLOAD_URL)**/"
            + self.collection.meta.onprem_provider_path
        )
        self.remote_blob_url = (
            "https://**REPLACE(DAAC_CLOUD_ALL_CONFIGREPO_HOST)**/**REPLACE(DAAC_CLOUD_ALL_CONFIGREPO_ORG)**/"
            + DEFAULT_CONFIG_TREE_NAME
            + "/-/blob/master/collections/"
            + self.collection.name
            + ".json"
        )
        self.cumulus_dashboard_collection_url = (
            self.target.dashboard
            + "/#/collections/collection/"
            + self.collection.name
            + "/"
            + self.collection.version
        )
        self.search_collection_url = (
            self.target.search + "/search?q=" + self.collection.name + "&ac=true"
        )
        self.search_granules_url = (
            self.target.search
            + "/search/granules/collection-details?p="
            + self.collection.cmr_id
            + "&ac=true"
        )
        self.cmr_collection_gem_url = (
            self.target.cmr + "/search/concepts/" + self.collection.cmr_id + ".html"
        )

    def write(self, instructions_writer):
        w = instructions_writer
        w.title(self.title)

        w.heading("Links", 1)
        for url in [
            self.jira_issue_url,
            self.onprem_details_url,
            self.onprem_download_url,
            self.remote_blob_url,
            self.cumulus_dashboard_collection_url,
            self.search_collection_url,
            self.search_granules_url,
            self.cmr_collection_gem_url,
        ]:
            w.unordered_list_item(url)
        w.close_list()

        w.heading("Onprem stats", 1)
        w.unordered_list_item("{:n} files".format(self.collection.onprem.file_count))
        w.unordered_list_item("{:n} bytes".format(self.collection.onprem.total_bytes))
        w.close_list()

        w.heading("Copy files to S3", 1)
        w.heading("Prepare AWS CLI", 2)
        w.step(
            "Add short-term credentials from CloudTamer",
            [r"# edit ~/.aws/credentials"],
            skip_session_placeholder=True,
        )
        w.step(
            "Set profile and Cumulus prefix",
            [
                CommandSetVariable(export=True, AWS_PROFILE=self.target.prefix),
                CommandSetVariable(export=True, CUMULUS_PREFIX=self.target.prefix),
            ],
            skip_session_placeholder=True,
        )
        w.step(
            "Gimmick for showing commands",
            [r'aws () { ( set -x ; command aws "$@" ) ; }'],
            skip_session_placeholder=True,
        )
        w.heading(self.collection.name, 2)
        w.step(
            "Set variables",
            [
                CommandSetVariable(SHORT_NAME=self.collection.name),
                CommandSetVariable(
                    ONPREM_PROVIDER_PATH=self.collection.meta.onprem_provider_path.rstrip(
                        "/"
                    )
                ),
                CommandSetVariable(
                    PRIVATE_PROVIDER_PATH=self.collection.meta.private_provider_path.rstrip(
                        "/"
                    )
                ),
                CommandSetVariable(EXAMPLE_FILENAME=self.collection.sample_file_name),
                r"SOURCE=/**REPLACE(DAAC_ONPREM_ALL_DATA_ROOT)**/**REPLACE(DAAC_ONPREM_LIVE_DATA_ENV_ABBREV)**/**REPLACE(DAAC_ONPREM_ALL_DATA_RELPATH)**/${ONPREM_PROVIDER_PATH}",
                r"DESTINATION=s3://${CUMULUS_PREFIX}-private/${PRIVATE_PROVIDER_PATH}",
                r"declare -p AWS_PROFILE CUMULUS_PREFIX SHORT_NAME ONPREM_PROVIDER_PATH PRIVATE_PROVIDER_PATH EXAMPLE_FILENAME SOURCE DESTINATION",
            ],
        )
        w.step(
            "Check source file count, total size, and example file",
            [
                r"find ${SOURCE}/ -type f | wc -l",
                r"du -hs ${SOURCE}/",
                r"find ${SOURCE}/ -name ${EXAMPLE_FILENAME} -ls",
            ],
        )
        w.step(
            "Check destination file count and total size before",
            [r"aws s3 ls --human --summarize --recursive ${DESTINATION}/ | tail -n2"],
            ignore_error=True,
        )
        w.step(
            "Dry-run sync",
            [
                r"time aws s3 sync --dryrun ${SOURCE}/ ${DESTINATION}/ | sed -e '2,$ s/^\((dryrun) upload: \).*$/\1/' | uniq -c"
            ],
        )
        w.step(
            "Copy files using batch queue (atd)",
            [
                r'echo "set -x ; time aws s3 sync --quiet ${SOURCE}/ ${DESTINATION}/ ; time aws s3 ls --human --summarize --recursive ${DESTINATION}/ | tail -n2 ; time aws s3 ls --recursive ${DESTINATION}/ | grep ${EXAMPLE_FILENAME}" | batch',
                r"atq",
            ],
        )
        w.step("Wait for email...", [])


def generate(
    jira_issue_key,
    dsshortname,
    target_name,
    config_tree_path,
    config_remote_url=None,
    pull=True,
    instructions_writer=None,
):
    instructions = generate_instructions(
        jira_issue_key,
        dsshortname,
        target_name,
        config_tree_path,
        config_remote_url,
        pull,
    )
    instructions.write(instructions_writer)


def generate_instructions(
    jira_issue_key,
    dsshortname,
    target_name,
    config_tree_path,
    config_remote_url=None,
    pull=True,
):
    target = get_target_by_name(target_name)
    tree = ConfigurationTree(config_tree_path)
    tree.prepare(config_remote_url, pull)
    collection_data = tree.load_collection_data(dsshortname)
    collection = Collection(collection_data)
    collection.assert_consistency()
    collection.ask_cmr_for_details(target)
    collection.ask_onprem_for_details(OnPremConfig(target_name))
    instructions = Instructions(collection, jira_issue_key, target)
    instructions.prepare()
    return instructions


DOCX_STYLE_NORMAL = "Normal"
DOCX_STYLE_CODE = "Procedure - Code"
DOCX_STYLE_COMMAND = "Procedure - Command"
DOCX_STYLE_COMMAND_SET_VARIABLE = DOCX_STYLE_COMMAND + " - Set Variable"


class InstructionsDocxWriter(InstructionsWriter):
    def __init__(self):
        starter = __file__[:-2] + "docx"
        self.document = docx.Document(starter)
        P = WD_STYLE_TYPE.PARAGRAPH
        s = self.document.styles

    def title(self, text):
        return self.document.add_heading(text, 0)

    def heading(self, text, level):
        return self.document.add_heading(text, level)

    def unordered_list_item(self, text):
        return self.document.add_paragraph(text, style="List Bullet")

    def step(
        self, description, commands, skip_session_placeholder=False, ignore_error=False
    ):
        self.document.add_paragraph("")
        if description:
            self.document.add_paragraph(description)
        for command in commands:
            if isinstance(command, basestring):
                command = Command(command)
            if isinstance(command, CommandSetVariable):
                text = ""
                if command.shell == "bash" and command.export:
                    text = "export "
                text += command.name
                text += "="
                p = self.document.add_paragraph(
                    text, style=DOCX_STYLE_COMMAND_SET_VARIABLE
                )
                r = p.add_run(command.value)
                r.bold = True
                r.italic = True
            elif isinstance(command, Command):
                self.document.add_paragraph(command.text, style=DOCX_STYLE_COMMAND)
        if not skip_session_placeholder:
            if commands:
                self.document.add_paragraph("")
            for _ in range(2):
                self.document.add_paragraph("", style=DOCX_STYLE_CODE)


class InstructionsHTMLWriter(InstructionsWriter):
    def __init__(self):
        self.html = ET.Element("html")
        self.stack = [ET.SubElement(self.html, "body")]

    def _cursor(self):
        return self.stack[-1]

    def _add_child(self, tag, push=False):
        elem = ET.SubElement(self._cursor(), tag)
        if push:
            self.stack.append(elem)
        return elem

    def title(self, text):
        return self.heading(text, 0)

    def heading(self, text, level):
        elem = self._add_child("h" + str(level + 1))
        elem.text = text
        return elem

    def _is_list(self, tag=None):
        tags = ("ul",)
        if tag is not None:
            tags = (tag,)
        return self._cursor().tag in tags

    def _start_list(self, tag):
        return self._add_child(tag, push=True)

    def _list_item(self, tag, text):
        if not self._is_list(tag):
            self._start_list(tag)
        return self._add_child("li", text)

    def unordered_list_item(self, text):
        return self._list_item("ul", text)

    def close_list(self):
        list_elem = None
        if self._is_list():
            list_elem = self._pop()
        return list_elem


class InstructionsBashWriter(InstructionsWriter):
    def __init__(self):
        self.lines = ["#!/usr/bin/bash", "set -euo pipefail", "set -v"]

    def _print(self, *values):
        sep = " "
        line = sep.join([str(_) for _ in values])
        self.lines.append(line)

    def title(self, text):
        bar = "#" * 60
        self._print(bar)
        self._print("#", text)
        self._print(bar)
        self._print()

    def heading(self, text, level):
        self._print()
        self._print("#" * (10 - level))
        self._print("#", text)
        self._print()

    def unordered_list_item(self, text):
        self._print("#  * ", text)

    def close_list(self):
        self._print("#")
        self._print()

    def step(
        self, description, commands, skip_session_placeholder=False, ignore_error=False
    ):
        step_lines = []
        if ignore_error:
            step_lines.append("set +e")
        for command in commands:
            if isinstance(command, basestring):
                command = Command(command)
            if isinstance(command, CommandSetVariable):
                text = ""
                if command.shell == "bash" and command.export:
                    text = "export "
                text += command.name
                text += "="
                text += command.value
                step_lines.append(text)
            elif isinstance(command, Command):
                step_lines.append(command.text)
        if description:
            self._print("#", description)
            self._print()
        for step_line in step_lines:
            self._print(step_line)
        if ignore_error:
            step_lines.append("set -e")
        self._print()


def generate_docx(
    jira_issue_key,
    dsshortname,
    target_name,
    config_tree_path,
    config_remote_url=None,
    pull=True,
):
    writer = InstructionsDocxWriter()
    instructions = generate_instructions(
        jira_issue_key=jira_issue_key,
        dsshortname=dsshortname,
        target_name=target_name,
        config_tree_path=config_tree_path,
        config_remote_url=config_remote_url,
        pull=pull,
    )
    instructions.write(writer)
    return writer.document


def generate_html(
    jira_issue_key,
    dsshortname,
    target_name,
    config_tree_path,
    config_remote_url=None,
    pull=True,
):
    writer = InstructionsHTMLWriter()
    instructions = generate_instructions(
        jira_issue_key=jira_issue_key,
        dsshortname=dsshortname,
        target_name=target_name,
        config_tree_path=config_tree_path,
        config_remote_url=config_remote_url,
        pull=pull,
    )
    instructions.write(writer)
    return ET.ElementTree(writer.html)


class ConfigurationTree:
    def __init__(self, tree, collections_path="collections"):
        self.tree = os.path.expanduser(tree)
        self.collections_path = collections_path

    def prepare(self, config_remote_url=None, pull=True):
        if not self.exists():
            self.clone(config_remote_url)
        elif pull:
            self.pull()

    def exists(self):
        return os.path.exists(self.tree)

    def clone(self, remote):
        subprocess.check_call(["git", "clone", remote, self.tree])

    def pull(self):
        subprocess.check_call(
            [
                "git",
                "=".join(("--work-tree", self.tree)),
                "=".join(("--git-dir", os.path.join(self.tree, ".git"))),
                "pull",
                "--ff-only",
            ]
        )

    def load_collection_data(self, dsshortname):
        json_path = os.path.join(
            self.tree, self.collections_path, ".".join((dsshortname, "json"))
        )
        json_file = open(json_path)
        collection_data = None
        with json_file:
            collection_data = json.load(json_file)
        name = collection_data["name"]
        assert name == dsshortname, (
            "'name' ("
            + repr(name)
            + ") does not match expected shortname ("
            + repr(dsshortname)
            + ")"
        )
        return collection_data


class CollectionMeta(dict):
    def __init__(self, data):
        super(CollectionMeta, self).__init__(data)
        self.private_provider_path = self["provider_path"]
        self.onprem_provider_path = self.private_provider_path

    def assert_consistency(self):
        assert hasattr(
            self, "private_provider_path"
        ), "'private_provider_path' is missing"


class CollectionOnPrem(dict):
    def __init__(self):
        super(CollectionOnPrem, self).__init__()
        self.config = None


class Collection(dict):
    def __init__(self, collection_data):
        super(Collection, self).__init__(collection_data)
        self.name = self["name"]
        files = self["files"]
        self.protected_sample_file_name = None
        for f in files:
            f_bucket = f.get("bucket")
            f_sample_file_name = f.get("sampleFileName")
            if f_bucket == "protected":
                self.protected_sample_file_name = f_sample_file_name
                break
        self.meta = CollectionMeta(self["meta"])
        parts = self.meta.private_provider_path.split("/")
        if self.name == parts[0]:
            self.meta.onprem_provider_path = "/".join(parts[1:])
        self.sample_file_name = self["sampleFileName"]
        self.url_path = self["url_path"]
        self.version = self["version"]
        self.onprem = CollectionOnPrem()

    def assert_consistency(self):
        self.meta.assert_consistency()
        expected_url_path = "__".join((self.name, self.version))
        assert self.url_path == expected_url_path, (
            "'url_path' ("
            + repr(self.url_path)
            + ") should use shortname and version ("
            + repr(expected_url_path)
            + ")"
        )
        beginnings = [self.name + "/", self.name + "_" + self.version + "/"]
        assert any(
            [self.meta.private_provider_path.startswith(_) for _ in beginnings]
        ), (
            "'meta.private_provider_path' ("
            + repr(self.meta.private_provider_path)
            + ") should start with one of: "
            + ", ".join(beginnings)
        )
        endings = ("/data/", "/browse/", "/%YYYY/%MM%DD/")
        assert any([self.meta.private_provider_path.endswith(_) for _ in endings]), (
            "'meta.private_provider_path' ("
            + repr(self.meta.private_provider_path)
            + ") should end with one of: "
            + ", ".join(endings)
        )

    def ask_cmr_for_details(self, target):
        found = False
        for cmr_preferred_provider in target.cmr_preferred_providers + [None]:
            cmr_search_url = target.cmr + "/search/collections"
            params = {"short_name": self.name, "version": self.version}
            if cmr_preferred_provider:
                params["provider"] = cmr_preferred_provider
            headers = {"accept": "application/json"}
            r = requests.get(cmr_search_url, params=params, headers=headers)
            r.raise_for_status()
            hits = r.headers["cmr-hits"]
            print(str(hits) + " hits from " + r.url, file=sys.stderr)
            if hits == "0":
                continue
            expected_hits = "1"
            assert hits == expected_hits, (
                "unexpected "
                + expected_hits
                + " CMR hit but saw "
                + hits
                + " ("
                + r.url
                + ")"
            )
            found = True
            data = r.json()
            feed = data["feed"]
            entry = feed["entry"][0]
            self.cmr_id = entry["id"]
            self.cmr_concept_id, self.cmr_provider_id = self.cmr_id.split("-", 1)
            # self.cmr_concept_id = "XXCONCEPT" + target.name.upper() + "XX"
            # self.cmr_provider_id = "XXPROVIDER" + target.name.upper() + "XX"
            # self.cmr_id = self.cmr_concept_id + "-" + self.cmr_provider_id
            break
        assert found, "no results found in CMR (" + r.url + ")"

    def ask_onprem_for_details(self, onprem_config):
        self.onprem.config = onprem_config
        self.onprem.source = os.path.join(
            self.onprem.config.source, self.meta.onprem_provider_path
        )
        file_count = 0
        total_bytes = 0
        empty_files = []
        for root, _, files in os.walk(self.onprem.source):
            for f in files:
                file_count += 1
                path = os.path.join(root, f)
                stats = os.stat(path)
                file_bytes = stats.st_size
                if file_bytes:
                    total_bytes += file_bytes
                else:
                    empty_files += path
        self.onprem.file_count = file_count
        self.onprem.total_bytes = total_bytes
        # self.onprem.empty_files = empty_files
        if empty_files:
            logging.warning("Empty files:")
            for f in sorted(empty_files):
                logging.warning(f)
            sys.exit(99)


class Command(object):
    def __init__(self, text):
        self.text = text


class CommandSetVariable(Command):
    def __init__(self, name=None, value=None, shell="bash", export=False, **kwargs):
        if name is None and value is None:
            name, value = kwargs.popitem()
        self.name = name
        self.value = value
        self.shell = shell
        self.export = export
        text = "=".join((self.name, self.value))
        if self.shell == "bash" and export:
            text = " ".join(("export", text))
        super(CommandSetVariable, self).__init__(text)


class MarkdownWriter:
    def __init__(self, out):
        self.out = out

    def _write(self, s):
        return self.out.write(str(s))

    def _newline(self):
        self._write(os.linesep)

    def heading(self, level, content):
        self._write(("#" * level) + " " + content)
        self._newline()
        self._newline()

    def unordered_list_item(self, content):
        self._write("* " + str(content))
        self._newline()


class MarkdownLink:
    def __init__(self, url, text=None):
        self.url = url
        if text is None:
            self.text = url

    def __str__(self):
        return "[" + self.text + "](" + self.url + ")"


class HTMLWriter:
    def __init__(self, out):
        self.out = out

    def _write(self, s):
        return self.out.write(str(s))

    def _newline(self):
        self._write(os.linesep)

    def heading(self, level, content):
        self._write(("#" * level) + " " + content)
        self._newline()
        self._newline()

    def unordered_list_item(self, content):
        self._write("* " + str(content))
        self._newline()


class HTMLLink:
    def __init__(self, url, text=None):
        self.url = url
        if text is None:
            self.text = url

    def __str__(self):
        return "[" + self.text + "](" + self.url + ")"


def init_parser(**kwargs):
    parser_options = PARSER_OPTIONS.copy()
    parser_options.update(kwargs)
    _treename = parser_options.pop("_treename")
    _remote = parser_options.pop("_remote")
    default_configtree = os.path.join("~", _treename)
    parser = argparse.ArgumentParser(**parser_options)
    parser.add_argument(
        "jira_issue_key",
        help="the JIRA key for affected issue (eg, **REPLACE(DAAC_CLOUD_JIRA_PROJECT_KEY)**-123)",
        metavar="JIRA-ISSUE",
    )
    parser.add_argument(
        "dsshortname",
        help="collection shortname",
        metavar="DSSHORTNAME",
    )
    parser.add_argument(
        "--config-tree",
        default=default_configtree,
        help="path to configuration tree (%(default)s)",
    )
    parser.add_argument(
        "--config-remote",
        default=_remote,
        help="git remote url for configuration tree (%(default)s)",
    )
    parser.add_argument(
        "--no-pull",
        action="store_false",
        dest="pull",
        help="do NOT pull latest configuration tree from remote",
    )
    parser.add_argument(
        "--no-bash-script",
        action="store_false",
        dest="create_bash_script",
        help="do NOT create the bash script",
    )
    parser.add_argument(
        "--no-email-docx",
        action="store_false",
        dest="email_docx",
        help="do NOT email the DOCX",
    )
    parser.add_argument(
        "--target",
        choices=[_.name for _ in KNOWN_TARGETS],
        default=DEFAULT_TARGET_NAME,
        help="target environment in **REPLACE(DAAC_ABBREV)** Cloud (%(default)s)",
    )
    return parser


def main():
    argv = sys.argv
    parser = init_parser(prog=argv[0])
    args = parser.parse_args(argv[1:])
    instructions = generate_instructions(
        args.jira_issue_key,
        args.dsshortname,
        args.target,
        args.config_tree,
        args.config_remote,
        pull=args.pull,
    )
    if args.create_bash_script:
        bash_writer = InstructionsBashWriter()
        instructions.write(bash_writer)
        script_name = ".".join(
            (
                instructions.title.lower().replace(" ", "-").replace(":", ""),
                str(int(time.time())),
                "bash",
            )
        )
        with open(script_name, "w") as out:
            for line in bash_writer.lines:
                print(line, file=out)
        print("Bash script saved to:", script_name)
    if args.email_docx:
        docx_writer = InstructionsDocxWriter()
        instructions.write(docx_writer)
        document = docx_writer.document
        msg = email.mime.multipart.MIMEMultipart()
        msg["From"] = SENDER
        msg["To"] = RECIPIENT
        msg["Subject"] = (
            "WIP: "
            + args.jira_issue_key
            + " "
            + args.dsshortname
            + " "
            + args.target
            + " "
            + NOW
        )
        body = "\n".join((" ".join(argv), HOSTNAME_SHORT, NOW))
        msg.attach(email.mime.text.MIMEText(body))
        part = email.mime.base.MIMEBase("application", "octet-stream")
        data = io.BytesIO()
        document.save(data)
        part.set_payload(data.getvalue())
        email.encoders.encode_base64(part)
        disposition_value = "; ".join(
            (
                "attachment",
                "=".join(
                    (
                        "filename",
                        ".".join(
                            (
                                " - ".join(
                                    (
                                        " ".join(
                                            (
                                                args.target,
                                                "static collections",
                                            )
                                        ),
                                        TODAY,
                                        " ".join(
                                            (
                                                args.jira_issue_key,
                                                args.dsshortname,
                                            )
                                        ),
                                    )
                                ),
                                "docx",
                            )
                        ).join('""'),
                    )
                ),
            )
        )
        part.add_header("Content-Disposition", disposition_value)
        msg.attach(part)
        smtp = smtplib.SMTP("localhost")
        smtp.sendmail(SENDER, RECIPIENT, msg.as_string())
        smtp.quit()
    return 0


if __name__ == "__main__":
    locale.setlocale(locale.LC_ALL, "")
    sys.exit(main())
